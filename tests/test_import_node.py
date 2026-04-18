"""Tests pour graph/nodes/import_node.py — Import de scenarios."""

import pytest
from graph.nodes.import_node import (
    import_scenario, ImportError as ImportNodeError,
    _parse_xlsx, _parse_markdown, _parse_csv, _parse_txt, _parse_docx, _parse_pdf,
    _split_long_paragraphs, _read_text_file,
)


class TestParseXlsx:
    """Tests du parsing Excel."""

    def test_basic_3_rows(self, sample_xlsx):
        steps = _parse_xlsx(sample_xlsx, sheet_name="PLAN")
        assert len(steps) == 3
        assert steps[0]["step_id"] == "1"
        assert steps[0]["text_original"] == "Bienvenue dans le portail."
        assert steps[2]["step_id"] == "3"

    def test_step_fields_complete(self, sample_xlsx):
        steps = _parse_xlsx(sample_xlsx, sheet_name="PLAN")
        for step in steps:
            assert "step_id" in step
            assert "text_original" in step
            assert step["text_tts"] == ""
            assert step["cleaning_status"] == "pending"
            assert step["speed_factor"] == 1.0

    def test_empty_cells_ignored(self, tmp_path):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "PLAN"
        ws.append(["Etape", "Scripts"])
        ws.append(["1", "Texte valide."])
        ws.append([None, "Texte sans id."])     # Pas de step_id
        ws.append(["3", None])                   # Pas de texte
        ws.append(["4", ""])                     # Texte vide apres strip
        ws.append(["5", "Autre texte valide."])
        path = tmp_path / "test_empty.xlsx"
        wb.save(str(path))

        steps = _parse_xlsx(str(path), sheet_name="PLAN")
        step_ids = [s["step_id"] for s in steps]
        assert "1" in step_ids
        assert "5" in step_ids
        # Les lignes sans step_id ou sans texte sont ignorees
        assert len(steps) == 2


class TestParseMarkdown:
    """Tests du parsing Markdown."""

    def test_basic_3_blocks(self, sample_markdown):
        steps = _parse_markdown(sample_markdown)
        assert len(steps) == 3
        assert steps[0]["step_id"] == "1"
        assert steps[0]["text_original"] == "Bienvenue dans le portail."
        assert steps[2]["step_id"] == "3"

    def test_empty_file(self, tmp_path):
        empty = tmp_path / "empty.md"
        empty.write_text("", encoding="utf-8")
        steps = _parse_markdown(str(empty))
        assert steps == []

    def test_no_matching_blocks(self, tmp_path):
        md = tmp_path / "no_steps.md"
        md.write_text("# Titre\n\nParagraphe sans etapes.", encoding="utf-8")
        steps = _parse_markdown(str(md))
        assert steps == []


class TestImportScenario:
    """Tests du noeud import_scenario."""

    def test_missing_file_raises_import_error(self):
        state = {
            "source_file": "/chemin/inexistant.xlsx",
            "excel_sheet": "PLAN",
            "steps": [],
            "iteration_count": 0,
        }
        with pytest.raises(ImportNodeError, match="introuvable"):
            import_scenario(state)

    def test_no_source_file_raises_import_error(self):
        state = {
            "source_file": "",
            "excel_sheet": "PLAN",
            "steps": [],
            "iteration_count": 0,
        }
        with pytest.raises(ImportNodeError, match="introuvable"):
            import_scenario(state)

    def test_xlsx_import_has_source_format(self, sample_xlsx):
        state = {
            "source_file": sample_xlsx,
            "excel_sheet": "PLAN",
            "steps": [],
            "iteration_count": 0,
        }
        result = import_scenario(state)
        assert result["source_format"] == "xlsx"
        assert len(result["steps"]) == 3
        assert result["iteration_count"] == 1

    def test_markdown_import_has_source_format(self, sample_markdown):
        state = {
            "source_file": sample_markdown,
            "excel_sheet": "PLAN",
            "steps": [],
            "iteration_count": 0,
        }
        result = import_scenario(state)
        assert result["source_format"] == "md"
        assert len(result["steps"]) == 3

    def test_incremental_import_no_duplicates(self, sample_xlsx):
        # Premier import
        state = {
            "source_file": sample_xlsx,
            "excel_sheet": "PLAN",
            "steps": [],
            "iteration_count": 0,
        }
        result1 = import_scenario(state)
        assert len(result1["steps"]) == 3

        # Deuxieme import avec les memes etapes deja presentes
        state2 = {
            "source_file": sample_xlsx,
            "excel_sheet": "PLAN",
            "steps": result1["steps"],
            "iteration_count": 0,
        }
        result2 = import_scenario(state2)
        # Pas de doublons : toujours 3 etapes
        assert len(result2["steps"]) == 3
        step_ids = [s["step_id"] for s in result2["steps"]]
        assert len(step_ids) == len(set(step_ids))


class TestParseCsv:
    """Tests du parsing CSV (PRD-016)."""

    def test_basic_csv(self, tmp_path):
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("id;texte\n1;Bonjour le monde.\n2;Au revoir.\n", encoding="utf-8")
        steps = _parse_csv(str(csv_file))
        assert len(steps) == 2
        assert steps[0]["step_id"] == "1"
        assert steps[0]["text_original"] == "Bonjour le monde."

    def test_csv_comma_separator(self, tmp_path):
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("id,texte\n1,Bonjour.\n2,Merci.\n", encoding="utf-8")
        steps = _parse_csv(str(csv_file))
        assert len(steps) == 2

    def test_csv_single_column_empty(self, tmp_path):
        csv_file = tmp_path / "single.csv"
        csv_file.write_text("texte\nBonjour\n", encoding="utf-8")
        steps = _parse_csv(str(csv_file))
        assert steps == []

    def test_csv_latin1(self, tmp_path):
        csv_file = tmp_path / "latin1.csv"
        csv_file.write_bytes("id;texte\n1;Bienvenue à l'accueil.\n".encode("latin-1"))
        steps = _parse_csv(str(csv_file))
        assert len(steps) == 1
        assert "accueil" in steps[0]["text_original"]


class TestParseTxt:
    """Tests du parsing TXT (PRD-016)."""

    def test_basic_paragraphs(self, tmp_path):
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Premier paragraphe.\n\nDeuxieme paragraphe.\n\nTroisieme.\n", encoding="utf-8")
        steps = _parse_txt(str(txt_file))
        assert len(steps) == 3
        assert steps[0]["step_id"] == "1"
        assert steps[0]["text_original"] == "Premier paragraphe."

    def test_empty_file(self, tmp_path):
        txt_file = tmp_path / "empty.txt"
        txt_file.write_text("", encoding="utf-8")
        steps = _parse_txt(str(txt_file))
        assert steps == []

    def test_long_paragraph_split(self, tmp_path):
        sentences = ["Ceci est la phrase numero %d du document." % i for i in range(25)]
        long_text = " ".join(sentences)
        txt_file = tmp_path / "long.txt"
        txt_file.write_text(long_text + "\n", encoding="utf-8")
        steps = _parse_txt(str(txt_file))
        assert len(steps) >= 2
        for s in steps:
            assert len(s["text_original"].split()) <= 110


class TestParseDocx:
    """Tests du parsing DOCX (PRD-016)."""

    def test_basic_docx(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_heading("Titre ignore", level=1)
        doc.add_paragraph("Premier paragraphe de texte.")
        doc.add_paragraph("Deuxieme paragraphe.")
        doc.add_heading("Autre titre ignore", level=2)
        doc.add_paragraph("Troisieme paragraphe.")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        steps = _parse_docx(str(path))
        assert len(steps) == 3
        assert steps[0]["text_original"] == "Premier paragraphe de texte."

    def test_headings_only_empty(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_heading("Titre 1", level=1)
        doc.add_heading("Titre 2", level=2)
        path = tmp_path / "headings.docx"
        doc.save(str(path))

        steps = _parse_docx(str(path))
        assert steps == []


class TestSplitLongParagraphs:
    """Tests du filet 100 mots."""

    def test_short_paragraph_unchanged(self):
        result = _split_long_paragraphs(["Phrase courte."])
        assert result == ["Phrase courte."]

    def test_long_paragraph_split(self):
        # Texte avec des phrases pour que le split fonctionne
        sentences = ["Ceci est une phrase numero %d." % i for i in range(30)]
        long = " ".join(sentences)  # ~210 mots, 30 phrases
        result = _split_long_paragraphs([long], max_words=100)
        assert len(result) >= 2
        for chunk in result:
            assert len(chunk.split()) <= 110

    def test_multiple_paragraphs_mixed(self):
        short = "Phrase courte."
        sentences = ["Voici la phrase %d du texte." % i for i in range(25)]
        long = " ".join(sentences)
        result = _split_long_paragraphs([short, long], max_words=100)
        assert len(result) >= 3
        assert result[0] == short


class TestReadTextFile:
    """Tests du fallback encodage."""

    def test_utf8(self, tmp_path):
        f = tmp_path / "utf8.txt"
        f.write_text("Bonjour à tous.", encoding="utf-8")
        assert "à" in _read_text_file(str(f))

    def test_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin1.txt"
        f.write_bytes("Bienvenue à l'accueil.".encode("latin-1"))
        assert "accueil" in _read_text_file(str(f))


class TestParsePdf:
    """Tests du parsing PDF (PRD-016 Phase 2)."""

    def test_basic_pdf(self, tmp_path):
        """PDF avec du texte -> etapes extraites."""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        c.drawString(72, 700, "Premier paragraphe du document.")
        c.drawString(72, 650, "")
        c.drawString(72, 600, "Deuxieme paragraphe important.")
        c.showPage()
        c.save()

        steps = _parse_pdf(str(pdf_path))
        assert len(steps) >= 1
        assert steps[0]["step_id"] == "1"

    def test_empty_pdf(self, tmp_path):
        """PDF sans texte -> liste vide."""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        pdf_path = tmp_path / "empty.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        c.showPage()
        c.save()

        steps = _parse_pdf(str(pdf_path))
        assert steps == []
